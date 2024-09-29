import unicodedata

from pathlib import Path
from typing import List, Optional
import pandas as pd
from llama_index.core.readers.base import BaseReader

from server.knowledge.base import Document


class DocxReader(BaseReader):
    """Read Docx files that respect table, using python-docx library

    Reader behavior:
        - All paragraphs are extracted as a Document
        - Each table ¬is extracted as a Document, rendered as a CSV string
        - The output is a list of Documents, concatenating the above
        (tables + paragraphs)
    """

    def load_data(
            self, file_path: Path, extra_info: Optional[dict] = None, **kwargs
    ) -> List[Document]:
        """Load data using Docx reader

        Args:
            file_path (Path): Path to .docx file

        Returns:
            List[Document]: list of documents extracted from the HTML file
        """
        import docx

        file_path = Path(file_path).resolve()

        doc = docx.Document(str(file_path))
        all_text = "\n".join([unicodedata.normalize("NFKC", p.text) for p in doc.paragraphs])
        pages = [all_text]

        tables = []
        for table in doc.tables:
            n_row = len(table.rows)
            n_col = len(table.columns)

            arrays = [["" for _ in range(n_row)] for _ in range(n_col)]

            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    arrays[j][i] = cell.text

            tables.append(pd.DataFrame({a[0]: a[1:] for a in arrays}))

        extra_info = extra_info or {}

        documents = [
            Document(
                text=table.to_csv(index=False).strip(),  # strip_special_chars_markdown()
                metadata={
                    "table_origin": table.to_csv(index=False), "type": "table", **extra_info,
                },
                metadata_template="",
                metadata_seperator="",
            )
            for table in tables
        ]

        documents.extend(
            [
                Document(text=non_table_text.strip(), metadata={"page_label": 1, **extra_info})
                for _, non_table_text in enumerate(pages)
            ]
        )

        return documents
